
from tkinter import *
import tkinter as tk
from pynput import mouse
from pynput.mouse import Button, Controller
import threading
import pyautogui
from PIL import Image, ImageTk  #pip install pillow
from tkinter import filedialog, messagebox
import time
import os
import docx
import pytesseract
import cv2  # if we want to recognize spreadsheet image, we need to install it by 'pip install opencv-python'

pytesseract.pytesseract.tesseract_cmd = r"C:\\Users\\f.he\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract.exe"
aimImG1 = './img/screenShot_tem1.png'
aimImG2 = './img/screenShot_tem2.png'
root = Tk()
# root.overrideredirect(True)
root.geometry("500x200+888+444")
# root.wm_attributes("-transparentcolor","white")
root.attributes("-topmost",True)

########################################################################


set1Lab = tk.Label(root, text="Set 1:")
left_Top_xLAB = tk.Label(root, text="left_Top_x:")
left_Top_yLAB = tk.Label(root, text="left_Top_y:")
leftBottom_yLAB = tk.Label(root, text="leftBottom_y:")
rightTop_xLAB = tk.Label(text="rightTop_x:")

set1Lab.grid(row=0, column=0)
left_Top_xLAB.grid(row=1, column=0)
left_Top_yLAB.grid(row=1, column=1)
leftBottom_yLAB.grid(row=1, column=2)
rightTop_xLAB.grid(row=1, column=3)


left_Top_xVal = tk.StringVar()
left_Top_yVal = tk.StringVar()
leftBottom_yVal = tk.StringVar()
rightTop_xVal = tk.StringVar()
left_Top_xTXT = tk.Entry(textvariable=left_Top_xVal)
left_Top_yTXT = tk.Entry(textvariable=left_Top_yVal)
leftBottom_yTXT = tk.Entry(textvariable=leftBottom_yVal)
rightTop_xTXT = tk.Entry(textvariable=rightTop_xVal)

left_Top_xTXT.grid(row=2, column=0)
left_Top_yTXT.grid(row=2, column=1)
leftBottom_yTXT.grid(row=2, column=2)
rightTop_xTXT.grid(row=2, column=3)

########################################################################


set1Lab_2 = tk.Label(root, text="Set 2:")
left_Top_xLAB_2 = tk.Label(root, text="left_Top_x:")
left_Top_yLAB_2 = tk.Label(root, text="left_Top_y:")
leftBottom_yLAB_2 = tk.Label(root, text="leftBottom_y:")
rightTop_xLAB_2 = tk.Label(text="rightTop_x:")

set1Lab_2.grid(row=3, column=0)
left_Top_xLAB_2.grid(row=4, column=0)
left_Top_yLAB_2.grid(row=4, column=1)
leftBottom_yLAB_2.grid(row=4, column=2)
rightTop_xLAB_2.grid(row=4, column=3)


left_Top_xVal_2 = tk.StringVar()
left_Top_yVal_2 = tk.StringVar()
leftBottom_yVal_2 = tk.StringVar()
rightTop_xVal_2 = tk.StringVar()
left_Top_xTXT_2 = tk.Entry(textvariable = left_Top_xVal_2)
left_Top_yTXT_2 = tk.Entry(textvariable = left_Top_yVal_2)
leftBottom_yTXT_2 = tk.Entry(textvariable = leftBottom_yVal_2)
rightTop_xTXT_2 = tk.Entry(textvariable = rightTop_xVal_2)

left_Top_xTXT_2.grid(row=5, column=0)
left_Top_yTXT_2.grid(row=5, column=1)
leftBottom_yTXT_2.grid(row=5, column=2)
rightTop_xTXT_2.grid(row=5, column=3)

###########################################################################################################
startBtn_recording = tk.Button(root, text = "START Recording!", command=lambda: btnClick('recording'))
startBtn_recording.grid(row=6, column=0)

startBtn_start = tk.Button(root, text = "START", command=lambda: btnClick('execute'))
startBtn_start.grid(row=7, column=0)

progressEntryVal = tk.StringVar()
progressEntry = tk.Entry(root, textvariable = progressEntryVal)
progressEntry.grid(row=7, column=1)




changePageClickLAB = tk.Label(root, text="Change Button:")

changePage_x_Val = tk.StringVar()
changePage_y_Val = tk.StringVar()


changePage_x = tk.Entry(textvariable = changePage_x_Val)
changePage_y = tk.Entry(textvariable = changePage_y_Val)

changePageClickLAB.grid(row=8, column=0)
changePage_x.grid(row=8, column=1)
changePage_y.grid(row=8, column=2)
#*******************************************************************
def btnClick(action):
    m1 =  MouseActionListener()
    m2 = MouseActionExecute()


    if action == "recording":
        m1.start()
    elif action == "execute":
        m2.start()


def initialSetting():
    progressEntryVal.set("0%")
    with open('./docss/location.txt') as file:
        setList = file.readlines()
        for i in range(0, len(setList)):
            it = setList[i]
            if "\\n" in it:
                newit = it.replace("\\n", "")
                setList[i] = newit
            if i == 0:
                smallList = list(it.strip().split(","))
                left_Top_xVal.set(smallList[0])
                left_Top_yVal.set(smallList[1])
                leftBottom_yVal.set(smallList[2])
                rightTop_xVal.set(smallList[3])
            elif i == 1:
                smallList = list(it.strip().split(","))
                left_Top_xVal_2.set(smallList[0])
                left_Top_yVal_2.set(smallList[1])
                leftBottom_yVal_2.set(smallList[2])
                rightTop_xVal_2.set(smallList[3])
            elif i ==2:
                smallList = list(it.strip().split(","))
                changePage_x_Val.set(smallList[0])
                changePage_y_Val.set(smallList[1])
        file.close()
#*********************************************************************
class MouseActionListener(threading.Thread):
    def __init__(self) -> None:
        super().__init__()
        self.x_changed = 0
        self.y_changed = 0


    def run(self):
        def on_move(x, y):
            """x, y: is the absolute axis number of mouse location"""
            # print("move: " + "x--" + str(x))
            # print("move: " + "y--" + str(y))
            pass

        def on_click(x, y, button, pressed):
            """
            x, y: is the absolute axis number of mouse click;
            button: is the identifer of mouse button(left/right/....);
            pressed: is a bool parameter, if true means pressed, if flase means release;
            """
            print("click: " + "x--" + str(x), str(button), str(pressed))
            print("click: " + "y--" + str(y), str(button), str(pressed))
            pass
        
        def on_scroll(x, y, x_axis, y_axis):
            """1)minus is scroll backwards;
               2)plus is scroll forwards
            if the changed number become 0, that means the screen recover to the original position.   
            """
            # self.x_changed += x_axis  
            # self.y_changed += y_axis  

            # print("click: " + "x: " + str(x) + "; x_changed: " + str(self.x_changed))
            # print("click: " + "y: " + str(y) + "; y_changed: " + str(self.y_changed))
            pass

        with mouse.Listener(on_move=on_move, on_click=on_click, on_scroll=on_scroll) as listener:
            listener.join()
        # with mouse.Listener(on_click=on_click) as listener:
        #     listener.join()


class MouseActionExecute(threading.Thread):
    def __init__(self) -> None:
        super().__init__()

    def run(self):
        mouse_exec = Controller()
        doc_obj = docx.Document()
        times_num = 0
        times_numLIMIT = 150

        #set11111111111111111
        left_int = int(left_Top_xTXT.get())
        top_int = int(left_Top_yTXT.get())
        width_int = int(rightTop_xTXT.get()) - left_int
        height_int = int(leftBottom_yTXT.get()) - top_int

        #set222222222222222222
        left_int_2 = int(left_Top_xTXT_2.get())
        top_int_2 = int(left_Top_yTXT_2.get())
        width_int_2 = int(rightTop_xTXT_2.get()) - left_int_2
        height_int_2 = int(leftBottom_yTXT_2.get()) - top_int_2

        while times_num < times_numLIMIT:

            #control mouse--------------------------------------------------------------------
            mouse_exec.position = [changePage_x_Val.get(), changePage_y_Val.get()] #!!!change the position of mouse directly
            mouse_exec.click(Button.left)
            
            #other event>>>>>>>>>>
            # mouse_exec.move()
            # mouse_exec.scroll()
            # mouse_exec.press()
            # mouse_exec.release()
            #<<<<<<<<<<<<<<<<<<<<<

            time.sleep(0.5)


            #screenshot--------------------------------------------------------------------
            #set11111111111111111
            part_ss = pyautogui.screenshot(region=(left_int, top_int, width_int , height_int))

            part_ss.save(r'./img/screenShot_tem1.png')   #  save the screenshot as a temporary picture for using later (convert that into the data can be saved into clipboard)
            # doc_obj.add_picture('./img/screenShot_tem.png')
            print(pytesseract.image_to_string(Image.open(aimImG1)))
            os.remove('./img/screenShot_tem.png')

            #set222222222222222222
            # part_ss_2 = pyautogui.screenshot(region=(left_int_2, top_int_2, width_int_2 , height_int_2))

            # part_ss_2.save(r'./img/screenShot_tem2.png')   #  save the screenshot as a temporary picture for using later (convert that into the data can be saved into clipboard)
            # # doc_obj.add_picture('./img/screenShot_tem.png')
            # print(pytesseract.image_to_string(Image.open(aimImG2)))
            # os.remove('./img/screenShot_tem.png')
            


            times_num += 1

            percetProgress = (times_num/times_numLIMIT) * 100
            progressEntryVal.set("%.2f" % (percetProgress,) + "%")

            time.sleep(3)

            
            


        # save docx
        doc_obj.save('./docss/ltoDOC.docx')
        messagebox.showinfo("Successfully", "Finished")

initialSetting()
root.mainloop()